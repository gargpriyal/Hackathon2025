from typing import Any
from fastapi import FastAPI, HTTPException, Depends, UploadFile, File, Form
from pymongo import AsyncMongoClient
from db import get_connection, close_connection
from models import User, Space, Chat, Topic, LevelOfUnderstanding, Document
from pydantic import BaseModel
from typing import Optional
from bson import ObjectId
import PyPDF2
from docx import Document as DocxDocument
import io
import voyageai
from datetime import datetime, timezone


class AddTopicToChat(BaseModel):
    user_id: str
    chat_id: str
    name: str


class UpdateTopicLevelOfUnderstanding(BaseModel):
    user_id: str
    name: str
    level_of_understanding: LevelOfUnderstanding


class SearchDocumentsRequest(BaseModel):
    query: str
    user_id: Optional[str] = None
    project_id: Optional[str] = None
    chat_id: Optional[str] = None
    limit: int = 5


model = "voyage-3-large"
vo = voyageai.Client()


def get_embedding(data, input_type="document"):
    embeddings = vo.embed(data, model=model, input_type=input_type).embeddings
    return embeddings[0]


async def get_db():
    client = get_connection()
    db = client['aivy_db']
    try:
        yield db
    finally:
        await close_connection(client)


app = FastAPI()


@app.get("/health")
def check_health():
    return {"status": "ok"}


@app.get("/db_status")
async def check_db_status(db: AsyncMongoClient = Depends(get_db)):
    return {"status": "ok"}


# ===== USER ENDPOINTS =====
@app.post("/users")
async def create_user(user: User, db: AsyncMongoClient = Depends(get_db)):
    collection = db['users']
    if await collection.find_one({"email": user.email}):
        raise HTTPException(status_code=400, detail="User already exists")
    created_user = await collection.insert_one(user.model_dump())
    return {"status": "success", "user_id": str(created_user.inserted_id)}


@app.get("/users")
async def get_users(db: AsyncMongoClient = Depends(get_db)):
    collection = db['users']
    users = await collection.find({}).to_list()
    for user in users:
        user["_id"] = str(user["_id"])
    return users


@app.get("/users/{user_id}")
async def get_user(user_id: str, db: AsyncMongoClient = Depends(get_db)):
    collection = db['users']
    user = await collection.find_one({"_id": ObjectId(user_id)})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    user["_id"] = str(user["_id"])
    return user


# ===== PROJECT ENDPOINTS =====
@app.post("/spaces")
async def create_project(project: Space, db: AsyncMongoClient = Depends(get_db)):
    collection = db['spaces']
    user_collection = db['users']
    user = await user_collection.find_one({"_id": ObjectId(project.user_id)})
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    if await collection.find_one({"project_name": project.project_name, "user_id": project.user_id}):
        raise HTTPException(status_code=400, detail="Space already exists")
    created_project = await collection.insert_one(project.model_dump())
    return {"status": "success", "project_id": str(created_project.inserted_id)}


@app.get("/spaces")
async def get_spaces(db: AsyncMongoClient = Depends(get_db)):
    collection = db['spaces']
    spaces = await collection.find().to_list()
    for project in spaces:
        project["_id"] = str(project["_id"])
    return spaces


@app.get("/spaces/user/{user_id}")
async def get_spaces_user(user_id: str, db: AsyncMongoClient = Depends(get_db)):
    collection = db['spaces']
    spaces = await collection.find({"user_id": user_id}).to_list()
    for project in spaces:
        project["_id"] = str(project["_id"])
    return spaces


@app.get("/spaces/{project_id}")
async def get_project(project_id: str, db: AsyncMongoClient = Depends(get_db)):
    collection = db['spaces']
    project = await collection.find_one({"_id": ObjectId(project_id)})
    if not project:
        raise HTTPException(status_code=404, detail="Space not found")
    project["_id"] = str(project["_id"])
    return project


@app.put("/spaces/{project_id}")
async def update_project(project_id: str, project_name: str, db: AsyncMongoClient = Depends(get_db)):
    collection = db['spaces']
    updated_project = await collection.update_one(
        {"_id": ObjectId(project_id)}, 
        {"$set": {"project_name": project_name}}
    )
    if updated_project.modified_count == 0:
        raise HTTPException(status_code=404, detail="Space not found")
    return {"status": "success", "project_id": str(project_id)}


@app.delete("/spaces/{project_id}")
async def delete_project(project_id: str, db: AsyncMongoClient = Depends(get_db)):
    collection = db['spaces']
    deleted_project = await collection.delete_one({"_id": ObjectId(project_id)})
    if deleted_project.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Space not found")
    return {"status": "success", "project_id": str(project_id)}


# ===== CHAT ENDPOINTS (UPDATED) =====
@app.post("/chats")
async def create_chat(chat: Chat, db: AsyncMongoClient = Depends(get_db)):
    """
    Create a new chat. Messages are now stored in checkpointer, not here.
    """
    collection = db['chats']
    project_collection = db['spaces']
    
    project = await project_collection.find_one({"_id": ObjectId(chat.project_id)})
    if not project:
        raise HTTPException(status_code=404, detail="Space not found")
    
    created_chat = await collection.insert_one(chat.model_dump())
    return {"status": "success", "chat_id": str(created_chat.inserted_id)}


@app.get("/chats")
async def get_chats(db: AsyncMongoClient = Depends(get_db)):
    collection = db['chats']
    chats = await collection.find().to_list()
    for chat in chats:
        chat["_id"] = str(chat["_id"])
    return chats


@app.get("/chats/project/{project_id}")
async def get_chats_project(project_id: str, db: AsyncMongoClient = Depends(get_db)):
    collection = db['chats']
    chats = await collection.find({"project_id": project_id}).to_list()
    for chat in chats:
        chat["_id"] = str(chat["_id"])
    return chats


@app.get("/chats/user/{user_id}")
async def get_chats_user(user_id: str, db: AsyncMongoClient = Depends(get_db)):
    collection = db['chats']
    chats = await collection.find({"user_id": user_id}).to_list()
    for chat in chats:
        chat["_id"] = str(chat["_id"])
    return chats


@app.get("/chats/{chat_id}")
async def get_chat(chat_id: str, db: AsyncMongoClient = Depends(get_db)):
    """
    Get chat metadata. For messages, query the checkpointer.
    """
    collection = db['chats']
    chat = await collection.find_one({"_id": ObjectId(chat_id)})
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")
    chat["_id"] = str(chat["_id"])
    return chat


@app.get("/chats/{chat_id}/messages")
async def get_chat_messages(chat_id: str, db: AsyncMongoClient = Depends(get_db)):
    """
    Get messages from the checkpointer for a specific chat.
    """
    # Verify chat exists
    chat_collection = db['chats']
    chat = await chat_collection.find_one({"_id": ObjectId(chat_id)})
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")
    
    # Query checkpointer
    checkpoints_collection = db['checkpoints']
    
    # Get the latest checkpoint for this thread
    checkpoint = await checkpoints_collection.find_one(
        {"thread_id": chat_id},
        sort=[("checkpoint_id", -1)]
    )
    
    if not checkpoint:
        return {"status": "success", "messages": []}
    
    messages = checkpoint.get("checkpoint", {}).get("channel_values", {}).get("messages", [])
    
    return {
        "status": "success",
        "messages": messages,
        "chat_id": chat_id
    }


@app.put("/chats/{chat_id}/metadata")
async def update_chat_metadata(
    chat_id: str, 
    metadata: dict, 
    db: AsyncMongoClient = Depends(get_db)
):
    """
    Update chat metadata (summary, topics, tags, etc.)
    """
    collection = db['chats']
    updated_chat = await collection.update_one(
        {"_id": ObjectId(chat_id)},
        {
            "$set": {
                "metadata": metadata,
                "last_updated": datetime.now(timezone.utc)
            }
        }
    )
    if updated_chat.modified_count == 0:
        raise HTTPException(status_code=404, detail="Chat not found")
    return {"status": "success", "chat_id": chat_id}


# DEPRECATED - Messages now handled by checkpointer
# Keep for backward compatibility if needed
@app.put("/bulk/messages/{chat_id}")
async def bulk_update_messages(
    chat_id: str, 
    bulk_update_messages: list[Any], 
    db: AsyncMongoClient = Depends(get_db)
):
    """
    DEPRECATED: Messages are now stored in the checkpointer.
    This endpoint is kept for backward compatibility.
    """
    raise HTTPException(
        status_code=410,
        detail="This endpoint is deprecated. Messages are now managed by the LangGraph checkpointer."
    )


# ===== TOPIC ENDPOINTS =====
@app.post("/add_topic_to_chat")
async def add_topic_to_chat(add_topic_to_chat: AddTopicToChat, db: AsyncMongoClient = Depends(get_db)):
    collection = db['topics']
    existing_topic = await collection.find_one({
        "name": add_topic_to_chat.name,
        "user_id": add_topic_to_chat.user_id
    })
    
    if not existing_topic:
        new_topic = Topic(
            user_id=add_topic_to_chat.user_id,
            name=add_topic_to_chat.name,
            related_chats=[add_topic_to_chat.chat_id],
            level_of_understanding=LevelOfUnderstanding.Learning
        )
        await collection.insert_one(new_topic.model_dump())
        return {"status": "success", "topic_name": new_topic.name}
    else:
        # Only add chat_id if not already in the list
        updated_chat = await collection.update_one(
            {
                "name": add_topic_to_chat.name,
                "user_id": add_topic_to_chat.user_id
            },
            {"$addToSet": {"related_chats": add_topic_to_chat.chat_id}}
        )
        return {"status": "success", "topic_name": add_topic_to_chat.name}


@app.put("/update_topic_level_of_understanding")
async def update_topic_level_of_understanding(
    update_topic: UpdateTopicLevelOfUnderstanding, 
    db: AsyncMongoClient = Depends(get_db)
):
    collection = db['topics']
    updated_topic = await collection.update_one(
        {
            "name": update_topic.name,
            "user_id": update_topic.user_id
        },
        {"$set": {"level_of_understanding": update_topic.level_of_understanding}}
    )
    if updated_topic.modified_count == 0:
        raise HTTPException(status_code=404, detail="Topic not found")
    return {"status": "success", "topic_name": update_topic.name}


@app.get("/topics")
async def get_topics(db: AsyncMongoClient = Depends(get_db)):
    collection = db['topics']
    topics = await collection.find().to_list()
    for topic in topics:
        topic["_id"] = str(topic["_id"])
    return topics


@app.get("/topics/user/{user_id}")
async def get_user_topics(user_id: str, db: AsyncMongoClient = Depends(get_db)):
    """
    Get all topics for a specific user
    """
    collection = db['topics']
    topics = await collection.find({"user_id": user_id}).to_list()
    for topic in topics:
        topic["_id"] = str(topic["_id"])
    return topics


@app.get("/topics/{topic_name}")
async def get_topic(topic_name: str, user_id: str, db: AsyncMongoClient = Depends(get_db)):
    """
    Get a topic from the database for a specific user
    """
    collection = db['topics']
    topic = await collection.find_one({
        "name": topic_name,
        "user_id": user_id
    })
    if not topic:
        raise HTTPException(status_code=404, detail="Topic not found")
    topic["_id"] = str(topic["_id"])
    return topic


@app.get("/topic_chats/{topic_name}")
async def get_chats_topic(
    topic_name: str, 
    user_id: str, 
    db: AsyncMongoClient = Depends(get_db)
):
    collection = db['topics']
    collection_chats = db['chats']
    
    topic = await collection.find_one({
        "name": topic_name,
        "user_id": user_id
    })
    if not topic:
        raise HTTPException(status_code=404, detail="Topic not found")
    
    topic = Topic(**topic)
    chats = []
    for chat_id in topic.related_chats:
        chat = await collection_chats.find_one({"_id": ObjectId(chat_id)})
        if chat:
            chat["_id"] = str(chat["_id"])
            chats.append(chat)
    
    return {"status": "success", "chats": chats}


# ===== DOCUMENT ENDPOINTS (UPDATED) =====
@app.post("/documents")
async def parse_file(
    user_id: str = Form(...),
    project_id: str = Form(...),
    chat_id: Optional[str] = Form(None),
    file: UploadFile = File(...),
    db=Depends(get_db),
):
    """
    Parse a file and add it to the database with a vector embedding
    """
    collection = db['documents']
    if not file:
        raise HTTPException(status_code=400, detail="File is required")
    
    file_content = await file.read()
    filename = file.filename.lower() if file.filename else ""
    
    if filename.endswith('.pdf'):
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text_content = ""
        for idx, page in enumerate(pdf_reader.pages):
            text_content += f"Page {idx + 1}\n"
            text_content += page.extract_text() + "\n"
        
        embedding = get_embedding(text_content)
        document = Document(
            user_id=user_id,
            project_id=project_id,
            name=filename,
            chat_id=chat_id,
            text_content=text_content.strip(),
            embedding=embedding,
            source="upload"
        )
        inserted_document = await collection.insert_one(document.model_dump())
        return {"status": "success", "document_id": str(inserted_document.inserted_id)}
        
    elif filename.endswith('.docx'):
        docx_file = io.BytesIO(file_content)
        doc = DocxDocument(docx_file)
        text_content = ""
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        embedding = get_embedding(text_content)
        document = Document(
            user_id=user_id,
            project_id=project_id,
            name=filename,
            chat_id=chat_id,
            text_content=text_content.strip(),
            embedding=embedding,
            source="upload"
        )
        inserted_document = await collection.insert_one(document.model_dump())
        return {"status": "success", "document_id": str(inserted_document.inserted_id)}
        
    elif filename.endswith(('.txt', '.md')):
        text_content = file_content.decode('utf-8')
        embedding = get_embedding(text_content)
        document = Document(
            user_id=user_id,
            project_id=project_id,
            name=filename,
            chat_id=chat_id,
            text_content=text_content.strip(),
            embedding=embedding,
            source="upload"
        )
        inserted_document = await collection.insert_one(document.model_dump())
        return {"status": "success", "document_id": str(inserted_document.inserted_id)}
        
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Supported formats: PDF, DOCX, TXT, MD"
        )


@app.get("/documents/{document_id}")
async def get_document(document_id: str, db: AsyncMongoClient = Depends(get_db)):
    collection = db['documents']
    document = await collection.find_one({"_id": ObjectId(document_id)})
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    document["_id"] = str(document["_id"])
    return document


@app.get("/documents")
async def get_documents(db: AsyncMongoClient = Depends(get_db)):
    collection = db['documents']
    documents = await collection.find().to_list()
    for document in documents:
        document["_id"] = str(document["_id"])
    return documents


@app.delete("/documents/{document_id}")
async def delete_document(document_id: str, db: AsyncMongoClient = Depends(get_db)):
    collection = db['documents']
    deleted_document = await collection.delete_one({"_id": ObjectId(document_id)})
    if deleted_document.deleted_count == 0:
        raise HTTPException(status_code=404, detail="Document not found")
    return {"status": "success", "document_id": str(document_id)}


@app.post("/search_documents")
async def search_documents(
    search_request: SearchDocumentsRequest,
    db: AsyncMongoClient = Depends(get_db)
):
    """
    Vector search for documents with filtering by user, project, and chat
    """
    collection = db['documents']
    query_embedding = get_embedding(search_request.query, input_type="query")
    
    # Build filter based on provided parameters
    filters = []
    
    # Priority: chat > project > user
    if search_request.chat_id:
        filters.append({"chat_id": search_request.chat_id})
    
    if search_request.project_id:
        filters.append({
            "project_id": search_request.project_id,
            "chat_id": None  # Space-level docs only
        })
    
    if search_request.user_id:
        # Include user's docs from other spaces
        user_filter = {"user_id": search_request.user_id}
        if search_request.project_id:
            user_filter["project_id"] = {"$ne": search_request.project_id}
        filters.append(user_filter)
    
    # If no filters provided, search all documents
    vector_search_filter = {"$or": filters} if filters else {}
    
    pipeline = [
        {
            "$vectorSearch": {
                "index": "vector_index",
                "queryVector": query_embedding,
                "path": "embedding",
                "exact": True,
                "limit": search_request.limit,
                **({"filter": vector_search_filter} if filters else {})
            }
        },
        {
            "$addFields": {
                "score": {"$meta": "vectorSearchScore"},
                # Boost chat-specific documents
                "relevance": {
                    "$cond": {
                        "if": {"$eq": ["$chat_id", search_request.chat_id]},
                        "then": {"$multiply": [{"$meta": "vectorSearchScore"}, 1.5]},
                        "else": {"$meta": "vectorSearchScore"}
                    }
                }
            }
        },
        {"$sort": {"relevance": -1}},
        {
            "$project": {
                "_id": {"$toString": "$_id"},
                "name": 1,
                "text_content": 1,
                "user_id": 1,
                "project_id": 1,
                "chat_id": 1,
                "source": 1,
                "score": 1,
                "relevance": 1
            }
        }
    ]
    
    results = await collection.aggregate(pipeline).to_list()
    
    return {
        "status": "success",
        "results": results,
        "count": len(results)
    }